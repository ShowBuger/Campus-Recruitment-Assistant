"""Multi-provider resume and job analysis — per-user isolation."""
import json
import os
import re
from pathlib import Path
from datetime import datetime
from uuid import uuid4
from urllib.parse import quote

import bleach
import markdown
import requests
from docx import Document
from fastapi import APIRouter, Depends, HTTPException, Response
from pydantic import BaseModel, Field
from pypdf import PdfReader

from app import ai_provider_utils, auth as auth_module, company_enrichment, database, local_records, state


router = APIRouter(prefix="/api/ai", tags=["ai"])
PROMPT_PATH = Path(__file__).resolve().parents[1] / "prompts" / "interview_analysis.md"
PROJECT_DIR = Path(__file__).resolve().parents[2]
ENRICHMENT_SKILL_PATH = PROJECT_DIR / "skills" / "company-job-enrichment" / "SKILL.md"
MAX_RESUME_CHARS = 40_000
PROVIDER_NAMES = {
    "deepseek": "DeepSeek",
    "openai": "OpenAI GPT",
    "anthropic": "Claude",
    "apidock": "ApiDock",
}
MARKDOWN_TAGS = {
    "h1", "h2", "h3", "h4", "h5", "h6", "p", "br", "hr",
    "ul", "ol", "li", "strong", "em", "blockquote", "code", "pre",
    "table", "thead", "tbody", "tr", "th", "td", "a",
}


class AnalysisRequest(BaseModel):
    resume_filename: str
    record_id: str
    analysis_mode: str = "match"
    focus: str = Field(default="", max_length=1000)


ANALYSIS_MODES = {
    "match": {
        "label": "综合匹配分析",
        "instruction": """请输出：
# 匹配结论
- 综合匹配度（0-100 分）、一句话判断、最适合强调的候选人定位
## JD 核心要求
## 简历匹配优势
列出 3-5 条，每条引用具体证据。
## 缺口与风险
列出 3-5 条并标注“未体现/证据不足/可能缺乏”。
## 简历修改建议
## 面试重点
## 针对性面试题
生成 8 道题，每题包含考察点、回答思路和可能追问。
## 7 天准备计划""",
    },
    "technical": {
        "label": "技术面试训练",
        "instruction": """生成 10 道针对性技术面试题：技术基础 3-4 道、项目深挖 3-4 道、系统/场景题 1-2 道、反问环节 1 道。
每题严格包含：题目类型、题目、依据（JD 或简历）、难度、参考回答框架、优秀回答应包含的证据、可能追问。
最后输出“最高风险知识点”和“面试前冲刺清单”。""",
    },
    "hr": {
        "label": "HR 面试训练",
        "instruction": """生成 8 道校招 HR 面专项问题，覆盖求职动机、公司与岗位理解、地点接受度、offer 选择、稳定性、协作冲突、职业规划和薪资预期。
每题包含：HR 真正判断什么、结合本简历的回答原则、可直接说出口的回答框架、踩坑提醒、可能追问。
最后总结候选人的 HR 面风险与应对策略。""",
    },
    "full": {
        "label": "完整面试流程",
        "instruction": """设计一套连贯的完整面试流程：
# 整体判断
## Round 1 · 一面
基础知识与简历入口，至少 5 题。
## Round 2 · 二面
项目深挖、技术方案和权衡，至少 5 题，并承接一面暴露的问题。
## Round 3 · 三面
综合判断、协作、业务理解和成长潜力，至少 4 题。
## Round 4 · HR 面
动机、稳定性、地点、薪资和规划，至少 5 题。
每轮说明关注点、通过关键、淘汰风险和准备建议。""",
    },
    "resume": {
        "label": "简历定向优化",
        "instruction": """针对目标岗位审查简历并输出：
# 简历诊断摘要
## 应保留和强化的内容
## 应删除或弱化的内容
## 缺失的关键词与证据
## 逐段修改建议
引用原表述，给出修改方向和示例表达；禁止虚构数据。
## 项目经历重写模板
## 面向该岗位的一分钟自我介绍
## 修改优先级清单
按 P0/P1/P2 排序。""",
    },
}


def _render_markdown(content: str) -> str:
    rendered = markdown.markdown(content, extensions=["fenced_code", "tables", "sane_lists"])
    return bleach.clean(
        rendered,
        tags=MARKDOWN_TAGS,
        attributes={"a": ["href", "title"]},
        protocols={"http", "https", "mailto"},
        strip=True,
    )


def _user_history_dir(user_id: int) -> Path:
    d = PROJECT_DIR / "data" / "users" / str(user_id) / "analysis_history"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _save_history(data: dict, user_id: int) -> str:
    hist_dir = _user_history_dir(user_id)
    history_id = datetime.now().strftime("%Y%m%d-%H%M%S") + "-" + uuid4().hex[:8]
    target = hist_dir / f"{history_id}.json"
    temp = target.with_suffix(".json.tmp")
    payload = {"id": history_id, "created_at": datetime.now().isoformat(timespec="seconds"), **data}
    temp.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    os.replace(temp, target)
    return history_id


def _history_path(history_id: str, user_id: int) -> Path:
    if not history_id or any(ch not in "0123456789abcdefghijklmnopqrstuvwxyz-" for ch in history_id.lower()):
        raise HTTPException(status_code=400, detail="无效的分析记录 ID")
    return _user_history_dir(user_id) / f"{history_id}.json"


def _user_resume_path(user_id: int, filename: str) -> Path:
    clean = Path(filename).name
    from app.routers.resume import ALLOWED_SUFFIXES
    if not clean or Path(clean).suffix.lower() not in ALLOWED_SUFFIXES:
        raise HTTPException(status_code=400, detail="无效的简历文件名")
    return PROJECT_DIR / "data" / "users" / str(user_id) / "resumes" / clean


@router.get("/history")
def list_history(user: dict = Depends(auth_module.get_current_user)):
    hist_dir = _user_history_dir(user["user_id"])
    items = []
    for path in sorted(hist_dir.glob("*.json"), reverse=True):
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
            items.append({
                key: data.get(key, "")
                for key in (
                    "id", "created_at", "company", "job", "resume", "provider", "model",
                    "analysis_mode", "analysis_mode_label",
                )
            })
        except (OSError, ValueError):
            continue
    return {"items": items}


@router.get("/history/{history_id}")
def get_history(history_id: str, user: dict = Depends(auth_module.get_current_user)):
    path = _history_path(history_id, user["user_id"])
    if not path.is_file():
        raise HTTPException(status_code=404, detail="分析记录不存在")
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, ValueError) as exc:
        raise HTTPException(status_code=422, detail="分析记录文件损坏") from exc
    analysis = str(data.get("analysis") or "")
    return {**data, "analysis_html": _render_markdown(analysis)}


@router.delete("/history/{history_id}")
def delete_history(history_id: str, user: dict = Depends(auth_module.get_current_user)):
    path = _history_path(history_id, user["user_id"])
    if not path.is_file():
        raise HTTPException(status_code=404, detail="分析记录不存在")
    try:
        path.unlink()
    except OSError as exc:
        raise HTTPException(status_code=500, detail="分析记录删除失败") from exc
    return {"success": True, "message": "分析记录已删除"}


@router.get("/history/{history_id}/download")
def download_history(history_id: str, user: dict = Depends(auth_module.get_current_user)):
    path = _history_path(history_id, user["user_id"])
    if not path.is_file():
        raise HTTPException(status_code=404, detail="分析记录不存在")
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, ValueError) as exc:
        raise HTTPException(status_code=422, detail="分析记录文件损坏") from exc
    analysis = str(data.get("analysis") or "")
    company = re.sub(r'[\\/:*?"<>|]+', "_", str(data.get("company") or "公司"))
    job = re.sub(r'[\\/:*?"<>|]+', "_", str(data.get("job") or "岗位"))
    created = str(data.get("created_at") or history_id).replace(":", "-").replace("T", "_")
    filename = f"{company}-{job}-{created}.md"
    return Response(
        content=analysis.encode("utf-8"),
        media_type="text/markdown; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"},
    )


def _extract_resume_text(path: Path) -> str:
    if path.suffix.lower() == ".pdf":
        reader = PdfReader(path)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    else:
        document = Document(path)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        text = "\n".join(parts)
    text = text.strip()
    if not text:
        raise HTTPException(status_code=422, detail="简历未提取到文本，扫描版 PDF 请先进行 OCR")
    return text[:MAX_RESUME_CHARS]


def _find_record(record_id: str, user_id: int) -> dict:
    if not record_id.startswith("rec"):
        raise HTTPException(status_code=422, detail="无效的本地记录 ID")
    record = local_records.get_record(user_id, record_id)
    if not record:
        raise HTTPException(status_code=404, detail="未找到所选总表记录")
    return record.get("fields") or {}


def _response_error(response: requests.Response) -> str:
    try:
        payload = response.json()
        error = payload.get("error") if isinstance(payload, dict) else None
        if isinstance(error, dict):
            return str(error.get("message") or error.get("type") or response.text)[:500]
        if error:
            return str(error)[:500]
        return str(payload)[:500]
    except Exception:
        return (response.text or f"HTTP {response.status_code}")[:500]


def _openai_output_text(payload: dict) -> str:
    parts = []
    for item in payload.get("output") or []:
        if not isinstance(item, dict) or item.get("type") != "message":
            continue
        for content in item.get("content") or []:
            if isinstance(content, dict) and content.get("type") == "output_text":
                parts.append(str(content.get("text") or ""))
    text = "\n".join(part for part in parts if part).strip()
    if not text:
        raise ValueError("OpenAI 响应中没有文本内容")
    return text


def _call_ai_provider(
    provider: str,
    api_key: str,
    model: str,
    system_prompt: str,
    user_content: str,
    base_url: str = "",
    api_mode: str = "responses",
    max_output_tokens: int = 12_000,
) -> str:
    safe_base = ai_provider_utils.validate_public_base_url(base_url, provider)
    # apidock uses OpenAI-compatible API, always chat_completions
    if provider == "apidock":
        api_mode = "chat_completions"
    if provider == "openai" or provider == "apidock":
        if api_mode == "chat_completions":
            response = requests.post(
                ai_provider_utils.endpoint_url(safe_base, "chat/completions"),
                headers=ai_provider_utils.auth_headers(provider, api_key),
                json={
                    "model": model,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_content},
                    ],
                    "stream": False,
                },
                timeout=180,
                allow_redirects=False,
            )
            if not response.ok:
                raise HTTPException(status_code=502, detail=f"OpenAI 兼容接口请求失败：{_response_error(response)}")
            return str(response.json()["choices"][0]["message"]["content"]).strip()
        response = requests.post(
            ai_provider_utils.endpoint_url(safe_base, "responses"),
            headers=ai_provider_utils.auth_headers(provider, api_key),
            json={
                "model": model,
                "instructions": system_prompt,
                "input": user_content,
                "max_output_tokens": max_output_tokens,
            },
            timeout=180,
            allow_redirects=False,
        )
        if not response.ok:
            raise HTTPException(status_code=502, detail=f"OpenAI API 请求失败：{_response_error(response)}")
        return _openai_output_text(response.json())

    if provider == "anthropic":
        response = requests.post(
            ai_provider_utils.endpoint_url(safe_base, "messages"),
            headers=ai_provider_utils.auth_headers(provider, api_key),
            json={
                "model": model,
                "max_tokens": max_output_tokens,
                "system": system_prompt,
                "messages": [{"role": "user", "content": user_content}],
            },
            timeout=180,
            allow_redirects=False,
        )
        if not response.ok:
            raise HTTPException(status_code=502, detail=f"Claude API 请求失败：{_response_error(response)}")
        payload = response.json()
        text = "\n".join(
            str(block.get("text") or "")
            for block in payload.get("content") or []
            if isinstance(block, dict) and block.get("type") == "text"
        ).strip()
        if not text:
            raise ValueError("Claude 响应中没有文本内容")
        return text

    response = requests.post(
        ai_provider_utils.endpoint_url(safe_base, "chat/completions"),
        headers=ai_provider_utils.auth_headers(provider, api_key),
        json={
            "model": model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content},
            ],
            "stream": False,
        },
        timeout=180,
        allow_redirects=False,
    )
    if not response.ok:
        raise HTTPException(status_code=502, detail=f"DeepSeek API 请求失败：{_response_error(response)}")
    return str(response.json()["choices"][0]["message"]["content"]).strip()


@router.post("/records/{record_id}/enrich")
def enrich_record(
    record_id: str,
    user: dict = Depends(auth_module.get_current_user),
):
    if not record_id.startswith("rec"):
        raise HTTPException(status_code=422, detail="无效的本地记录 ID")
    record = local_records.get_record(user["user_id"], record_id)
    if not record:
        raise HTTPException(status_code=404, detail="未找到对应的个人总表记录")
    fields = record["fields"]
    company = str(fields.get("公司名称") or "").strip()
    job = str(fields.get("秋招岗位") or "").strip()
    if not company or not job:
        raise HTTPException(status_code=422, detail="AI 补全需要先填写并保存公司名称和岗位")

    cfg = database.get_user_config(user["user_id"])
    provider = cfg.get("ai_provider") or "deepseek"
    if provider not in PROVIDER_NAMES:
        provider = "deepseek"
    key_field = {
        "deepseek": "deepseek_api_key",
        "openai": "openai_api_key",
        "anthropic": "anthropic_api_key",
    }[provider]
    model_field = {
        "deepseek": "deepseek_model",
        "openai": "openai_model",
        "anthropic": "anthropic_model",
    }[provider]
    model_default = {
        "deepseek": "deepseek-v4-flash",
        "openai": "gpt-5.4-mini",
        "anthropic": "claude-sonnet-5",
        "apidock": "gpt-4o",
    }[provider]
    api_key = cfg.get(key_field, "")
    if not api_key:
        raise HTTPException(status_code=400, detail=f"请先在 AI 配置中填写 {PROVIDER_NAMES[provider]} API Key")
    model = cfg.get(model_field, "") or model_default
    base_url = cfg.get(f"{provider}_base_url", "") or ai_provider_utils.DEFAULT_BASE_URLS[provider]
    api_mode = cfg.get("openai_api_mode", "") or "responses"
        except company_enrichment.EnrichmentError:
            evidence = []
        current_type = (fields.get("公司/行业类型") or [""])[0]
        current_directions = fields.get("嵌入式方向") or []
        current_note = str(fields.get("备注") or "")
        knowledge_fallback = not evidence
        skill_prompt = ENRICHMENT_SKILL_PATH.read_text(encoding="utf-8")
        evidence_content = company_enrichment.evidence_text(evidence) if evidence else (
            "未取得网页证据，切换为模型知识模式。仅使用你有较高把握的、稳定的既有知识；"
            "不要声称已联网或已核实，不要编造来源。对于不熟悉的公司或不确定的信息，"
            "对应字段返回空值。note_append 应直接陈述有把握的公司业务和岗位背景，"
            "不要输出“根据名称推断”“无来源核实”“请以实际为准”等模板化免责声明。"
        )
        user_content = f"""请根据以下记录与联网搜索证据执行公司岗位信息补全。

【公司名称】
{company}

【岗位名称】
{job}

【已有公司类型】
{current_type or '空'}

【已有方向】
{'、'.join(current_directions) if current_directions else '空'}

【联网搜索证据】
{evidence_content}
"""
        raw_result = _call_ai_provider(
            provider, api_key, model, skill_prompt, user_content,
            base_url=base_url, api_mode=api_mode, max_output_tokens=1800,
        )
        enrichment = company_enrichment.parse_result(
            raw_result,
            evidence,
            allow_unsourced_note=knowledge_fallback,
            allow_empty=knowledge_fallback,
        )
        updates = {}
        updated_fields = []
        if not current_type and enrichment["company_type"]:
            updates["公司/行业类型"] = enrichment["company_type"]
            updated_fields.append("公司类型")
        if not current_directions and enrichment["directions"]:
            updates["嵌入式方向"] = enrichment["directions"]
            updated_fields.append("方向")
        merged_note = company_enrichment.appended_note(
            current_note, enrichment, datetime.now().strftime("%Y-%m-%d")
        )
        if merged_note != current_note:
            updates["备注"] = merged_note
            updated_fields.append("备注")
        if updates:
            local_records.update_record(user["user_id"], record_id, updates)
        dashboard = local_records.get_dashboard_data(user["user_id"])
        state.set_cache(user["user_id"], dashboard)
    except HTTPException:
        raise
    except company_enrichment.EnrichmentError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except requests.RequestException as exc:
        raise HTTPException(status_code=502, detail=f"无法连接 {PROVIDER_NAMES[provider]} API：{exc}") from exc
    except (OSError, KeyError, IndexError, TypeError, ValueError) as exc:
        raise HTTPException(status_code=502, detail=f"AI 补全返回格式异常：{exc}") from exc

    return {
        "success": True,
        "message": "AI 补全完成" if updated_fields else "没有需要补充的新内容",
        "updated_fields": updated_fields,
        "company_type": enrichment["company_type"] if not current_type else current_type,
        "directions": enrichment["directions"] if not current_directions else current_directions,
        "note": merged_note,
        "sources": enrichment["sources"],
        "provider": provider,
        "model": model,
        "dashboard": dashboard,
    }


@router.post("/analyze")
def analyze_resume(
    request: AnalysisRequest,
    user: dict = Depends(auth_module.get_current_user),
):
    cfg = database.get_user_config(user["user_id"])
    provider = cfg.get("ai_provider") or "deepseek"
    if provider not in PROVIDER_NAMES:
        provider = "deepseek"
    api_key_field = {
        "deepseek": "deepseek_api_key",
        "openai": "openai_api_key",
        "anthropic": "anthropic_api_key",
    }[provider]
    model_field = {
        "deepseek": "deepseek_model",
        "openai": "openai_model",
        "anthropic": "anthropic_model",
    }[provider]
    model_default = {
        "deepseek": "deepseek-v4-flash",
        "openai": "gpt-5.4-mini",
        "anthropic": "claude-sonnet-5",
        "apidock": "gpt-4o",
    }[provider]
    api_key = cfg.get(api_key_field, "")
    if not api_key:
        raise HTTPException(status_code=400, detail=f"请先在 AI 配置中填写 {PROVIDER_NAMES[provider]} API Key")
    model = cfg.get(model_field, "") or model_default
    base_url = cfg.get(f"{provider}_base_url", "") or ai_provider_utils.DEFAULT_BASE_URLS[provider]
    api_mode = cfg.get("openai_api_mode", "") or "responses"

    resume_path = _user_resume_path(user["user_id"], request.resume_filename)
    if not resume_path.is_file():
        raise HTTPException(status_code=404, detail="所选简历不存在")
    mode = ANALYSIS_MODES.get(request.analysis_mode)
    if not mode:
        raise HTTPException(status_code=422, detail="不支持的分析模式")

    fields = _find_record(request.record_id, user["user_id"])
    resume_text = _extract_resume_text(resume_path)
    system_prompt = PROMPT_PATH.read_text(encoding="utf-8")
    company = str(fields.get("公司名称") or "")
    job = str(fields.get("秋招岗位") or "")
    job_jd = str(fields.get("岗位JD") or "")
    if not job_jd.strip():
        raise HTTPException(status_code=422, detail="所选记录尚未填写岗位 JD")

    user_content = f"""请分析以下岗位与简历：

【公司名称】
{company}

【岗位名称】
{job}

【岗位 JD】
{job_jd}

【候选人简历】
{resume_text}

【本次任务】
{mode['label']}

【输出要求】
{mode['instruction']}

【用户特别关注】
{request.focus.strip() or '无，请按默认流程全面分析。'}
"""
    try:
        content = _call_ai_provider(
            provider, api_key, model, system_prompt, user_content,
            base_url=base_url, api_mode=api_mode,
        )
    except HTTPException:
        raise
    except requests.RequestException as exc:
        raise HTTPException(status_code=502, detail=f"无法连接 {PROVIDER_NAMES[provider]} API：{exc}") from exc
    except (KeyError, IndexError, TypeError, ValueError) as exc:
        raise HTTPException(status_code=502, detail=f"{PROVIDER_NAMES[provider]} API 返回格式异常：{exc}") from exc

    safe_html = _render_markdown(content)
    history_id = _save_history({
        "company": company,
        "job": job,
        "resume": resume_path.name,
        "provider": provider,
        "model": model,
        "analysis_mode": request.analysis_mode,
        "analysis_mode_label": mode["label"],
        "record_id": request.record_id,
        "analysis": content,
    }, user["user_id"])
    return {
        "success": True,
        "analysis": content,
        "analysis_html": safe_html,
        "model": model,
        "provider": provider,
        "provider_name": PROVIDER_NAMES[provider],
        "company": company,
        "job": job,
        "resume": resume_path.name,
        "history_id": history_id,
        "analysis_mode": request.analysis_mode,
        "analysis_mode_label": mode["label"],
    }
