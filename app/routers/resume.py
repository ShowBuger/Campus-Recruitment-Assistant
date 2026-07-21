"""Resume upload, listing, and preview — per-user isolation."""
import html
import os
import zipfile
from pathlib import Path

from docx import Document
from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from fastapi.responses import FileResponse, HTMLResponse

from app import auth as auth_module

router = APIRouter(prefix="/api/resumes", tags=["resumes"])
PROJECT_DIR = Path(__file__).resolve().parents[2]
ALLOWED_SUFFIXES = {".pdf", ".docx"}
MAX_FILE_SIZE = 20 * 1024 * 1024


def _user_resume_dir(user_id: int) -> Path:
    d = PROJECT_DIR / "data" / "users" / str(user_id) / "resumes"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _resume_path(user_id: int, filename: str) -> Path:
    clean = Path(filename).name
    if not clean or clean != filename or Path(clean).suffix.lower() not in ALLOWED_SUFFIXES:
        raise HTTPException(status_code=400, detail="无效的简历文件名")
    return _user_resume_dir(user_id) / clean


def _file_info(path: Path) -> dict:
    stat = path.stat()
    return {
        "name": path.name,
        "type": path.suffix.lower()[1:],
        "size": stat.st_size,
        "modified": int(stat.st_mtime * 1000),
    }


@router.get("")
def list_resumes(user: dict = Depends(auth_module.get_current_user)):
    resume_dir = _user_resume_dir(user["user_id"])
    files = [p for p in resume_dir.iterdir() if p.is_file() and p.suffix.lower() in ALLOWED_SUFFIXES]
    files.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    return {"files": [_file_info(path) for path in files]}


@router.delete("/{filename}")
def delete_resume(filename: str, user: dict = Depends(auth_module.get_current_user)):
    path = _resume_path(user["user_id"], filename)
    if not path.is_file():
        raise HTTPException(status_code=404, detail="简历文件不存在")
    try:
        path.unlink()
    except OSError as exc:
        raise HTTPException(status_code=500, detail="简历文件删除失败") from exc
    return {"success": True, "message": "简历已删除"}


@router.post("")
async def upload_resume(
    file: UploadFile = File(...),
    user: dict = Depends(auth_module.get_current_user),
):
    filename = Path(file.filename or "").name
    target = _resume_path(user["user_id"], filename)
    temp = target.with_suffix(target.suffix + ".uploading")
    size = 0
    try:
        with temp.open("wb") as output:
            while chunk := await file.read(1024 * 1024):
                size += len(chunk)
                if size > MAX_FILE_SIZE:
                    raise HTTPException(status_code=413, detail="简历文件不能超过 20 MB")
                output.write(chunk)
        if target.suffix.lower() == ".pdf":
            with temp.open("rb") as source:
                if source.read(5) != b"%PDF-":
                    raise HTTPException(status_code=400, detail="文件内容不是有效的 PDF")
        else:
            if not zipfile.is_zipfile(temp):
                raise HTTPException(status_code=400, detail="文件内容不是有效的 DOCX")
            with zipfile.ZipFile(temp) as archive:
                names = set(archive.namelist())
                if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                    raise HTTPException(status_code=400, detail="文件内容不是有效的 DOCX")
        os.replace(temp, target)
    finally:
        await file.close()
        if temp.exists():
            temp.unlink()
    return {"success": True, "message": "简历已上传", "file": _file_info(target)}


@router.get("/{filename}/file")
def get_resume_file(filename: str, user: dict = Depends(auth_module.get_current_user)):
    path = _resume_path(user["user_id"], filename)
    if not path.is_file():
        raise HTTPException(status_code=404, detail="简历文件不存在")
    media_type = "application/pdf" if path.suffix.lower() == ".pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return FileResponse(path, media_type=media_type, content_disposition_type="inline", filename=path.name)


def _runs_html(paragraph) -> str:
    parts = []
    for run in paragraph.runs:
        value = html.escape(run.text).replace("\n", "<br>")
        if not value:
            continue
        if run.bold:
            value = f"<strong>{value}</strong>"
        if run.italic:
            value = f"<em>{value}</em>"
        if run.underline:
            value = f"<u>{value}</u>"
        parts.append(value)
    return "".join(parts) or html.escape(paragraph.text)


def _docx_html(path: Path) -> str:
    document = Document(path)
    body = []
    for block in document.element.body:
        tag = block.tag.rsplit("}", 1)[-1]
        if tag == "p":
            paragraph = next((p for p in document.paragraphs if p._p is block), None)
            if paragraph is None or not paragraph.text.strip():
                continue
            style = (paragraph.style.name or "").lower() if paragraph.style else ""
            level = "h2" if "heading 1" in style else "h3" if "heading 2" in style else "p"
            body.append(f"<{level}>{_runs_html(paragraph)}</{level}>")
        elif tag == "tbl":
            table = next((t for t in document.tables if t._tbl is block), None)
            if table is None:
                continue
            rows = []
            for row in table.rows:
                cells = "".join(f"<td>{html.escape(cell.text).replace(chr(10), '<br>')}</td>" for cell in row.cells)
                rows.append(f"<tr>{cells}</tr>")
            body.append(f"<table>{''.join(rows)}</table>")
    title = html.escape(path.name)
    return f"""<!doctype html><html lang="zh-CN"><head><meta charset="utf-8"><title>{title}</title>
<style>body{{margin:0;background:#eef2f8;color:#172033;font:15px/1.65 Arial,'Microsoft YaHei',sans-serif}}main{{width:min(900px,calc(100% - 32px));min-height:calc(100vh - 48px);margin:24px auto;padding:48px 56px;box-sizing:border-box;background:#fff;box-shadow:0 8px 30px rgba(15,23,42,.12)}}h2{{font-size:24px;margin:20px 0 8px}}h3{{font-size:18px;margin:18px 0 6px}}p{{margin:5px 0;white-space:pre-wrap}}table{{width:100%;border-collapse:collapse;margin:12px 0}}td{{padding:7px 9px;border:1px solid #dbe2ea;vertical-align:top}}@media(max-width:640px){{main{{margin:0;width:100%;padding:24px 18px;box-shadow:none}}}}</style>
</head><body><main>{''.join(body) or '<p>该 DOCX 没有可提取的文本内容。</p>'}</main></body></html>"""


@router.get("/{filename}/preview")
def preview_resume(filename: str, user: dict = Depends(auth_module.get_current_user)):
    path = _resume_path(user["user_id"], filename)
    if not path.is_file():
        raise HTTPException(status_code=404, detail="简历文件不存在")
    if path.suffix.lower() == ".pdf":
        return FileResponse(path, media_type="application/pdf", content_disposition_type="inline")
    try:
        return HTMLResponse(_docx_html(path))
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"DOCX 预览失败：{exc}") from exc
